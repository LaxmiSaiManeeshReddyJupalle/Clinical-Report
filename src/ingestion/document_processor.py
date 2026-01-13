"""
Document Processor Module for UIC ATU Clinical Report Generator

Extracts text content from various document formats (PDF, DOCX, TXT, RTF).
Provides chunking strategies for RAG pipeline integration.

SECURITY: Processes documents containing PHI. All operations are logged
without exposing document content.
"""

import logging
from pathlib import Path
from typing import List, Optional, Dict, Any, Union
from dataclasses import dataclass
from enum import Enum
import io

# Document parsing libraries
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# Configure logging (no PHI in logs)
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types."""
    PDF = ".pdf"
    DOCX = ".docx"
    DOC = ".doc"
    TXT = ".txt"
    RTF = ".rtf"
    UNKNOWN = ""


@dataclass
class DocumentChunk:
    """Represents a chunk of document text for embedding."""
    text: str
    metadata: Dict[str, Any]
    chunk_index: int
    start_char: int
    end_char: int

    @property
    def char_count(self) -> int:
        return len(self.text)


@dataclass
class ProcessedDocument:
    """Result of document processing."""
    source_path: str
    document_type: DocumentType
    total_chars: int
    total_pages: Optional[int]
    raw_text: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    success: bool
    error_message: Optional[str] = None


class TextExtractor:
    """
    Extracts text content from various document formats.

    Supported formats:
    - PDF (.pdf)
    - Word (.docx)
    - Plain text (.txt)
    - Rich text (.rtf) - basic support

    SECURITY: Document content is extracted but never logged.
    """

    @staticmethod
    def detect_type(file_path: Union[str, Path]) -> DocumentType:
        """
        Detect document type from file extension.

        Args:
            file_path: Path to the document

        Returns:
            DocumentType enum value
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        for doc_type in DocumentType:
            if doc_type.value == ext:
                return doc_type

        return DocumentType.UNKNOWN

    @staticmethod
    def extract_from_pdf(content: bytes) -> tuple[str, int]:
        """
        Extract text from PDF document.

        Args:
            content: PDF file content as bytes

        Returns:
            Tuple of (extracted text, page count)
        """
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            num_pages = len(reader.pages)

            text_parts = []
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted text from PDF: {num_pages} pages")
            return full_text, num_pages

        except Exception as e:
            logger.error(f"PDF extraction failed: {type(e).__name__}")
            raise

    @staticmethod
    def extract_from_docx(content: bytes) -> tuple[str, int]:
        """
        Extract text from Word document.

        Args:
            content: DOCX file content as bytes

        Returns:
            Tuple of (extracted text, paragraph count as pseudo-pages)
        """
        try:
            docx_file = io.BytesIO(content)
            doc = DocxDocument(docx_file)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            full_text = "\n\n".join(text_parts)
            # Use paragraph count as pseudo-page count
            page_count = max(1, len(text_parts) // 20)

            logger.info(f"Extracted text from DOCX: {len(text_parts)} paragraphs")
            return full_text, page_count

        except Exception as e:
            logger.error(f"DOCX extraction failed: {type(e).__name__}")
            raise

    @staticmethod
    def extract_from_txt(content: bytes) -> tuple[str, int]:
        """
        Extract text from plain text file.

        Args:
            content: Text file content as bytes

        Returns:
            Tuple of (text content, 1 for page count)
        """
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                text = content.decode('latin-1')

            logger.info("Extracted text from TXT file")
            return text, 1

        except Exception as e:
            logger.error(f"TXT extraction failed: {type(e).__name__}")
            raise

    @staticmethod
    def extract_from_rtf(content: bytes) -> tuple[str, int]:
        """
        Extract text from RTF file (basic support).

        Args:
            content: RTF file content as bytes

        Returns:
            Tuple of (text content, 1 for page count)
        """
        try:
            # Basic RTF parsing - strip RTF control codes
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                text = content.decode('latin-1')

            # Very basic RTF stripping (removes common control sequences)
            import re
            # Remove RTF header
            text = re.sub(r'\\rtf\d+.*?\\viewkind\d+', '', text)
            # Remove control words
            text = re.sub(r'\\[a-z]+\d*\s?', '', text)
            # Remove braces
            text = re.sub(r'[{}]', '', text)
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text).strip()

            logger.info("Extracted text from RTF file (basic parsing)")
            return text, 1

        except Exception as e:
            logger.error(f"RTF extraction failed: {type(e).__name__}")
            raise

    def extract(self, content: bytes, doc_type: DocumentType) -> tuple[str, int]:
        """
        Extract text from document based on type.

        Args:
            content: Document content as bytes
            doc_type: Type of document

        Returns:
            Tuple of (extracted text, page/section count)
        """
        extractors = {
            DocumentType.PDF: self.extract_from_pdf,
            DocumentType.DOCX: self.extract_from_docx,
            DocumentType.DOC: self.extract_from_docx,  # Try DOCX parser for DOC
            DocumentType.TXT: self.extract_from_txt,
            DocumentType.RTF: self.extract_from_rtf,
        }

        extractor = extractors.get(doc_type)
        if not extractor:
            raise ValueError(f"Unsupported document type: {doc_type}")

        return extractor(content)


class TextChunker:
    """
    Splits text into chunks suitable for embedding and retrieval.

    Supports multiple chunking strategies:
    - Fixed size with overlap
    - Sentence-based
    - Paragraph-based
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        min_chunk_size: int = 100
    ):
        """
        Initialize chunker with configuration.

        Args:
            chunk_size: Target size for each chunk in characters
            chunk_overlap: Number of overlapping characters between chunks
            min_chunk_size: Minimum chunk size (smaller chunks are merged)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size

    def chunk_by_size(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[DocumentChunk]:
        """
        Split text into fixed-size chunks with overlap.

        Args:
            text: Text to chunk
            metadata: Base metadata to include in each chunk

        Returns:
            List of DocumentChunk objects
        """
        if not text or not text.strip():
            return []

        chunks = []
        start = 0
        chunk_index = 0
        text_length = len(text)
        base_metadata = metadata or {}

        while start < text_length:
            # Calculate end position
            end = min(start + self.chunk_size, text_length)

            # Try to break at sentence boundary if not at end
            if end < text_length:
                # Look for sentence endings near the boundary
                search_start = max(start + self.min_chunk_size, end - 100)
                sentence_ends = [
                    text.rfind('. ', search_start, end),
                    text.rfind('.\n', search_start, end),
                    text.rfind('? ', search_start, end),
                    text.rfind('! ', search_start, end),
                ]
                best_end = max(sentence_ends)
                if best_end > search_start:
                    end = best_end + 1  # Include the period

            chunk_text = text[start:end].strip()

            if len(chunk_text) >= self.min_chunk_size:
                chunk_metadata = {
                    **base_metadata,
                    'chunk_index': chunk_index,
                    'start_char': start,
                    'end_char': end
                }

                chunks.append(DocumentChunk(
                    text=chunk_text,
                    metadata=chunk_metadata,
                    chunk_index=chunk_index,
                    start_char=start,
                    end_char=end
                ))
                chunk_index += 1

            # Move start position with overlap
            start = end - self.chunk_overlap if end < text_length else text_length

        logger.info(f"Created {len(chunks)} chunks from text")
        return chunks

    def chunk_by_paragraph(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[DocumentChunk]:
        """
        Split text by paragraphs, merging small ones.

        Args:
            text: Text to chunk
            metadata: Base metadata to include in each chunk

        Returns:
            List of DocumentChunk objects
        """
        if not text or not text.strip():
            return []

        # Split by double newlines (paragraphs)
        paragraphs = text.split('\n\n')
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        chunks = []
        current_chunk = []
        current_length = 0
        chunk_index = 0
        char_position = 0
        base_metadata = metadata or {}

        for para in paragraphs:
            para_length = len(para)

            # If adding this paragraph exceeds chunk size, save current chunk
            if current_length + para_length > self.chunk_size and current_chunk:
                chunk_text = '\n\n'.join(current_chunk)
                start_char = char_position - current_length
                chunks.append(DocumentChunk(
                    text=chunk_text,
                    metadata={
                        **base_metadata,
                        'chunk_index': chunk_index,
                        'start_char': start_char,
                        'end_char': char_position
                    },
                    chunk_index=chunk_index,
                    start_char=start_char,
                    end_char=char_position
                ))
                chunk_index += 1
                current_chunk = []
                current_length = 0

            current_chunk.append(para)
            current_length += para_length + 2  # +2 for \n\n
            char_position += para_length + 2

        # Don't forget the last chunk
        if current_chunk:
            chunk_text = '\n\n'.join(current_chunk)
            start_char = char_position - current_length
            chunks.append(DocumentChunk(
                text=chunk_text,
                metadata={
                    **base_metadata,
                    'chunk_index': chunk_index,
                    'start_char': start_char,
                    'end_char': char_position
                },
                chunk_index=chunk_index,
                start_char=start_char,
                end_char=char_position
            ))

        logger.info(f"Created {len(chunks)} paragraph-based chunks")
        return chunks


class DocumentProcessor:
    """
    Main document processing pipeline.

    Combines extraction and chunking for end-to-end document processing.

    SECURITY: Handles PHI documents - content is never logged.
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        chunking_strategy: str = "size"
    ):
        """
        Initialize document processor.

        Args:
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks
            chunking_strategy: "size" or "paragraph"
        """
        self.extractor = TextExtractor()
        self.chunker = TextChunker(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        self.chunking_strategy = chunking_strategy

    def process_file(
        self,
        file_path: Union[str, Path],
        content: Optional[bytes] = None,
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """
        Process a document file and return extracted, chunked content.

        Args:
            file_path: Path to the document (for metadata)
            content: Optional pre-loaded file content
            additional_metadata: Extra metadata to include

        Returns:
            ProcessedDocument with extracted text and chunks
        """
        path = Path(file_path)
        doc_type = self.extractor.detect_type(path)

        # Log without exposing file name (may contain PHI)
        logger.info(f"Processing document of type: {doc_type.value}")

        # Load content if not provided
        if content is None:
            try:
                content = path.read_bytes()
            except Exception as e:
                logger.error(f"Failed to read file: {type(e).__name__}")
                return ProcessedDocument(
                    source_path=str(path),
                    document_type=doc_type,
                    total_chars=0,
                    total_pages=0,
                    raw_text="",
                    chunks=[],
                    metadata={},
                    success=False,
                    error_message=f"Failed to read file: {type(e).__name__}"
                )

        try:
            # Extract text
            raw_text, page_count = self.extractor.extract(content, doc_type)

            # Build base metadata
            metadata = {
                'source': str(path.name),
                'document_type': doc_type.value,
                'page_count': page_count,
                'total_chars': len(raw_text),
                **(additional_metadata or {})
            }

            # Chunk the text
            if self.chunking_strategy == "paragraph":
                chunks = self.chunker.chunk_by_paragraph(raw_text, metadata)
            else:
                chunks = self.chunker.chunk_by_size(raw_text, metadata)

            return ProcessedDocument(
                source_path=str(path),
                document_type=doc_type,
                total_chars=len(raw_text),
                total_pages=page_count,
                raw_text=raw_text,
                chunks=chunks,
                metadata=metadata,
                success=True
            )

        except Exception as e:
            logger.error(f"Document processing failed: {type(e).__name__}")
            return ProcessedDocument(
                source_path=str(path),
                document_type=doc_type,
                total_chars=0,
                total_pages=0,
                raw_text="",
                chunks=[],
                metadata={},
                success=False,
                error_message=f"Processing failed: {type(e).__name__}"
            )

    def process_multiple(
        self,
        file_paths: List[Union[str, Path]],
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> List[ProcessedDocument]:
        """
        Process multiple documents.

        Args:
            file_paths: List of document paths
            additional_metadata: Extra metadata for all documents

        Returns:
            List of ProcessedDocument objects
        """
        results = []
        successful = 0
        failed = 0

        for file_path in file_paths:
            result = self.process_file(file_path, additional_metadata=additional_metadata)
            results.append(result)
            if result.success:
                successful += 1
            else:
                failed += 1

        logger.info(f"Processed {len(file_paths)} documents: {successful} successful, {failed} failed")
        return results


def create_document_processor(
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    strategy: str = "size"
) -> DocumentProcessor:
    """
    Factory function to create document processor.

    Args:
        chunk_size: Target chunk size
        chunk_overlap: Chunk overlap
        strategy: Chunking strategy ("size" or "paragraph")

    Returns:
        Configured DocumentProcessor
    """
    return DocumentProcessor(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        chunking_strategy=strategy
    )
